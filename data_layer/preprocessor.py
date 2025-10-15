import re
import json
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from sentence_transformers import SentenceTransformer
import numpy as np

# Import thêm cho xử lý file PDF, OCR, DOCX
from PyPDF2 import PdfReader
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
from docx import Document as DocxDocument
from pdf2image import convert_from_path
import pytesseract


class UXOPreprocessor:
    """
    Lớp tiền xử lý (preprocessor) cho dữ liệu UXO:
    - Làm sạch văn bản
    - Phân loại loại tài liệu (UXO info, safety, contact,...)
    - Chia nhỏ tài liệu (chunk)
    - Sinh embeddings để phục vụ RAG
    """
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        # Khởi tạo model SentenceTransformer để embedding văn bản
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)

    # -------------------------------
    # Làm sạch văn bản
    # -------------------------------
    def clean_text(self, text: str) -> str:
        """Loại bỏ HTML, ký tự rác, và các phần footer không cần thiết"""
        text = re.sub(r'<.*?>', '', text) # Xóa thẻ HTML
        text = re.sub(r'(javascript:|window\.|var\s+)', '', text, flags=re.IGNORECASE)
        # Loại bỏ các đoạn không liên quan thường xuất hiện cuối trang web
        noise_patterns = [
            r'Follow us.*', r'Subscribe.*', r'Contact us.*', r'©.*\d{4}.*',
            r'Terms of Use.*', r'Privacy Policy.*', r'All rights reserved.*',
            r'Sitemap.*', r'Search.*',
        ]
        for pat in noise_patterns:
            text = re.sub(pat, '', text, flags=re.IGNORECASE)
        # Thay thế ký tự HTML đặc biệt thành dạng đọc được
        html_entities = {
            '&nbsp;': ' ', '&amp;': '&', '&quot;': '"',
            '&apos;': "'", '&lt;': '<', '&gt;': '>',
            '\u2013': '-', '\u2014': '-', '\u2022': '•',
        }
        for k, v in html_entities.items():
            text = text.replace(k, v)
        # Chuẩn hóa khoảng trắng
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    # -------------------------------
    # Phân loại tài liệu
    # -------------------------------
    def process_documents(self, documents):
        """Làm sạch + phân loại tài liệu theo nội dung"""
        processed_docs = []
        for doc in documents:
            content = self.clean_text(doc.page_content)
            metadata = doc.metadata
            doc_type = "general" # Mặc định
            # Phân loại dựa trên từ khóa trong nội dung
            if re.search(r"\b(safety|an toàn|hướng dẫn)\b", content, re.IGNORECASE):
                doc_type = "safety_guidelines"
            elif re.search(r"\b(hotline|liên hệ|contact)\b", content, re.IGNORECASE):
                doc_type = "contact_info"
            elif re.search(r"\b(bom|mìn|uxo|ordnance)\b", content, re.IGNORECASE):
                doc_type = "uxo_info"
            # Tạo đối tượng Document mới với metadata đã gán
            processed_doc = Document(
                page_content=content,
                metadata={**metadata, "type": doc_type}
            )
            processed_docs.append(processed_doc)
        return processed_docs

    # -------------------------------
    # Cắt tài liệu thành chunk nhỏ
    # -------------------------------
    def split_documents(self, documents, chunk_size=1000, chunk_overlap=200):
        """
        Chia tài liệu dài thành các đoạn nhỏ (chunk)
        - chunk_size: độ dài mỗi đoạn
        - chunk_overlap: phần trùng giữa các đoạn (giúp mô hình không mất ngữ cảnh)
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        return splitter.split_documents(documents)

    # -------------------------------
    # Sinh vector embeddings
    # -------------------------------
    def embed_documents(self, documents):
        """Chuyển văn bản thành vector embeddings để lưu vào Chroma"""
        texts = [doc.page_content for doc in documents]
        embeddings = self.model.encode(
            texts, show_progress_bar=True, batch_size=32, normalize_embeddings=True
        )
        return embeddings

    # -------------------------------
    # Lưu dữ liệu đã xử lý ra file
    # -------------------------------
    def save_to_jsonl(self, documents, filename):
        """Lưu tài liệu sau khi làm sạch ra file JSONL"""
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        with open(filename, "w", encoding="utf-8") as f:
            for doc in documents:
                f.write(json.dumps({"content": doc.page_content, "metadata": doc.metadata}, ensure_ascii=False) + "\n")

    def save_embeddings(self, embeddings, documents, out_path="data/uxo_embeddings.npz"):
        """Lưu embeddings và metadata ra file .npz để dùng lại"""
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        np.savez_compressed(out_path, embeddings=embeddings, metadata=[doc.metadata for doc in documents])

    # -------------------------------
    # Pipeline tích hợp cho run.py
    # -------------------------------
    def clean_and_chunk(self, raw_docs, chunk_size=1000, chunk_overlap=200):
        """Làm sạch → phân loại → chunk văn bản"""
        processed = self.process_documents(raw_docs)
        chunks = self.split_documents(processed, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return chunks

    # -------------------------------
    # Đọc PDF (text hoặc OCR)
    # -------------------------------
    def read_pdf(self, file_path: str) -> str:
        """
        Đọc file PDF:
        - Ưu tiên đọc text trực tiếp
        - Nếu không có text (file scan), fallback sang OCR bằng Tesseract
        """
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Lỗi đọc PDF trực tiếp: {e}")

        # Nếu không đọc được text nào, thử OCR
        if not text.strip():
            if not OCR_AVAILABLE:
                print(f"Không thể đọc text và OCR không khả dụng: {file_path}")
                return ""
            try:
                # OCR bằng pdf2image + pytesseract
                poppler_path = r"E:\Poppler\poppler-24.07.0\Library\bin"  # Cập nhật đường dẫn poppler nếu cần
                images = convert_from_path(file_path, poppler_path=poppler_path)
                for i, img in enumerate(images):
                    ocr_text = pytesseract.image_to_string(img, lang='vie+eng')
                    text += ocr_text + "\n"
            except Exception as e:
                print(f"OCR thất bại cho file {file_path}: {e}")
                return ""

        return text

    # -------------------------------
    # Đọc file TXT
    # -------------------------------
    def read_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    # đọc DOCX
    def read_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
